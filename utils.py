import sounddevice as sd
from scipy.io.wavfile import write
import numpy as np
from docx import Document

def record_audio(duration, sample_rate=44100):
    print("Recording...")
    recording = sd.rec(int(duration * sample_rate), samplerate=sample_rate, channels=2, dtype='float64')
    sd.wait()  
    print("Recording finished")
    return recording, sample_rate

def save_as_wav(audio_data, sample_rate, filename):
    write(filename, sample_rate, audio_data)
    print(f"File saved as {filename}")


def find_and_insert_text(doc, title, text):
    """
    Finds a title in the document and inserts text immediately after the title paragraph.
    If the title is not found, it adds the title and text at the end of the document.
    """
    title_found = False
    for i, paragraph in enumerate(doc.paragraphs):
        if title_found:
            # Inserts text in the next paragraph space available after the title
            if i + 1 < len(doc.paragraphs):
                doc.paragraphs[i + 1].insert_paragraph_before(text)
            else:
                doc.add_paragraph(text)  
            return True
        if paragraph.text.lower() == title.lower():
            title_found = True
    doc.add_paragraph(text)  # Add the text
    return True